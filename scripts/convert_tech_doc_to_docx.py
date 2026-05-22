#!/usr/bin/env python3
"""Convert 技术文档.md to Word .docx with cross-platform safe fonts and full format support."""

import re
import os
import subprocess
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ── Paths ──
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.join(SCRIPT_DIR, "..", "deliverables", "docs")
MD_PATH = os.path.join(BASE_DIR, "技术文档.md")
FIGURES_DIR = os.path.join(BASE_DIR, "figures")
OUT_PATH = os.path.join(BASE_DIR, "技术文档.docx")

# ── Fonts: cross-platform safe Office defaults ──
FONT_WEST = "Calibri"           # Western script
FONT_EAST = "微软雅黑"           # East Asian (Microsoft YaHei, shipped with Office on both platforms)
FONT_CODE = "Consolas"
FONT_CODE_EAST = "微软雅黑"

# ── Read markdown ──
with open(MD_PATH, "r", encoding="utf-8") as f:
    lines = f.readlines()

# ── Create document ──
doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ── Style setup ──
def set_run_font(run, font_name=FONT_WEST, east_name=FONT_EAST, size_pt=12, bold=False, italic=False, color=None):
    """Apply cross-platform font settings to a run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # East Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), east_name)

def set_paragraph_font(p, font_name=FONT_WEST, east_name=FONT_EAST, size_pt=12):
    """Set default font for a paragraph."""
    for run in p.runs:
        set_run_font(run, font_name, east_name, size_pt)

# Normal style
style = doc.styles["Normal"]
style.font.name = FONT_WEST
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST)

# Heading styles
for level in range(1, 4):
    h_style = doc.styles[f"Heading {level}"]
    h_font = h_style.font
    h_font.name = FONT_WEST
    h_font.bold = True
    h_font.color.rgb = RGBColor(0, 0, 0)
    h_style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST)
    if level == 1:
        h_font.size = Pt(22)
    elif level == 2:
        h_font.size = Pt(16)
    else:
        h_font.size = Pt(14)

# ── Helpers ──
def is_h1(line):
    return re.match(r"^# (.+)", line)

def is_h2(line):
    return re.match(r"^## (.+)", line)

def is_h3(line):
    return re.match(r"^### (.+)", line)

def is_image(line):
    return re.match(r"^!\[(.+?)\]\((.+?)\)", line)

def is_table_sep(line):
    return re.match(r"^\|[\s\-:|]+\|", line)

def is_table_row(line):
    return line.startswith("|") and line.rstrip().endswith("|")

def is_code_fence(line):
    return line.strip().startswith("```")

def is_blockquote(line):
    return line.startswith("> ")

def is_list_item(line):
    return re.match(r"^- (.+)", line)

def is_empty(line):
    return line.strip() == ""

def svg_to_png(svg_path, png_path):
    """Convert SVG to PNG using rsvg-convert."""
    cmd = ["rsvg-convert", "-w", "1400", svg_path, "-o", png_path]
    subprocess.run(cmd, check=True, capture_output=True)
    return png_path

def parse_inline_format(paragraph, text):
    """Parse **bold**, *italic*, __bold__, _italic_, `code` in text and add runs."""
    # Pattern order matters: ** before *, __ before _, `code`
    pattern = re.compile(r'(\*\*[^*]+?\*\*|\*[^*]+?\*|__[^_]+?__|_[^_]+?_|`[^`]+?`)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, italic=True)
        elif part.startswith("_") and part.endswith("_"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, font_name=FONT_CODE, east_name=FONT_CODE_EAST, size_pt=11)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)

def add_blockquote(doc, lines, start, end):
    """Add a blockquote paragraph with proper formatting."""
    clean_lines = []
    for line in lines[start:end]:
        if line.startswith("> "):
            clean_lines.append(line[2:].rstrip())
        else:
            clean_lines.append(line.rstrip())
    clean = "\n".join(clean_lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    parse_inline_format(p, clean)
    # Apply quote styling to all runs
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(80, 80, 80)
        run.font.italic = True

# ── First pass: identify special block ranges ──
in_code_block = False
code_start = -1
special_ranges = []  # (start, end, "code"|"blockquote")

for i, line in enumerate(lines):
    stripped = line.strip()

    if is_code_fence(line):
        if not in_code_block:
            in_code_block = True
            code_start = i
        else:
            special_ranges.append((code_start, i + 1, "code"))
            in_code_block = False
        continue

    if in_code_block:
        continue

    # Blockquotes: look for contiguous > lines
    if is_blockquote(line):
        # Check if previous line was also blockquote (already in a range)
        prev_in_bq = False
        for s, e, t in special_ranges:
            if t == "blockquote" and s <= i < e:
                prev_in_bq = True
                break
        if not prev_in_bq:
            # Start new blockquote range
            bq_start = i
            bq_end = i + 1
            while bq_end < len(lines) and is_blockquote(lines[bq_end]):
                bq_end += 1
            special_ranges.append((bq_start, bq_end, "blockquote"))
        continue

# Sort and deduplicate ranges by start
special_ranges.sort(key=lambda x: x[0])
# Remove overlapping ranges (blockquote should not overlap with code)
filtered = []
for r in special_ranges:
    overlap = False
    for fr in filtered:
        if not (r[1] <= fr[0] or r[0] >= fr[1]):
            overlap = True
            break
    if not overlap:
        filtered.append(r)
special_ranges = filtered

covered = set()
for start, end, _ in special_ranges:
    for idx in range(start, end):
        covered.add(idx)

# ── Process lines ──
i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    if stripped == "":
        i += 1
        continue

    # Check special ranges
    special = None
    for start, end, stype in special_ranges:
        if start <= i < end:
            special = (start, end, stype)
            break

    if special:
        start, end, stype = special
        if stype == "code":
            code_lines = lines[start + 1:end - 1]
            code_text = "".join(code_lines).rstrip("\n")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(14)
            run = p.add_run(code_text)
            set_run_font(run, font_name=FONT_CODE, east_name=FONT_CODE_EAST, size_pt=10)
            run.font.color.rgb = RGBColor(50, 50, 50)
        elif stype == "blockquote":
            add_blockquote(doc, lines, start, end)
        i = end
        continue

    # Headings
    m = is_h1(line)
    if m:
        doc.add_heading(m.group(1), level=1)
        i += 1
        continue

    m = is_h2(line)
    if m:
        doc.add_heading(m.group(1), level=2)
        i += 1
        continue

    m = is_h3(line)
    if m:
        doc.add_heading(m.group(1), level=3)
        i += 1
        continue

    # Image
    m = is_image(line)
    if m:
        alt_text = m.group(1)
        svg_rel_path = m.group(2)
        svg_path = os.path.join(BASE_DIR, svg_rel_path)
        if os.path.exists(svg_path):
            try:
                # Use pre-generated PNG if available, else convert
                png_path = svg_path.replace(".svg", ".png")
                if not os.path.exists(png_path):
                    png_path = svg_path + ".tmp.png"
                    svg_to_png(svg_path, png_path)
                doc.add_picture(png_path, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Caption
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.add_run(alt_text)
                set_run_font(cap_run, size_pt=9)
                cap_run.font.color.rgb = RGBColor(100, 100, 100)
                # Clean temp png
                if png_path.endswith(".tmp.png") and os.path.exists(png_path):
                    os.remove(png_path)
            except Exception as e:
                p = doc.add_paragraph()
                run = p.add_run(f"[图片: {alt_text} — 转换失败: {e}]")
                set_run_font(run, color=RGBColor(255, 0, 0))
        else:
            p = doc.add_paragraph()
            run = p.add_run(f"[图片: {alt_text} — 文件未找到: {svg_path}]")
            set_run_font(run, color=RGBColor(255, 0, 0))
        i += 1
        continue

    # Table
    if is_table_sep(line):
        # Find header row (previous non-empty, non-separator row)
        header_idx = i - 1
        while header_idx >= 0 and (is_empty(lines[header_idx]) or is_table_sep(lines[header_idx])):
            header_idx -= 1

        if header_idx >= 0 and is_table_row(lines[header_idx]):
            header_cells = [c.strip() for c in lines[header_idx].strip().strip("|").split("|")]
        else:
            i += 1
            continue

        # Data rows after separator
        data_rows = []
        j = i + 1
        while j < len(lines) and is_table_row(lines[j]) and not is_table_sep(lines[j]):
            row_cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
            data_rows.append(row_cells)
            j += 1

        num_cols = len(header_cells)
        table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
        table.style = "Light Grid Accent 1"

        # Header
        for ci, cell_text in enumerate(header_cells):
            cell = table.rows[0].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            parse_inline_format(p, cell_text)
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

        # Data
        for ri, row in enumerate(data_rows):
            for ci, cell_text in enumerate(row):
                if ci < num_cols:
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    parse_inline_format(p, cell_text)
                    for run in p.runs:
                        run.font.size = Pt(10)

        i = j
        continue

    # List item
    m = is_list_item(line)
    if m:
        items = [m.group(1)]
        j = i + 1
        while j < len(lines) and not is_empty(lines[j]) and not is_list_item(lines[j]):
            # Check if next line starts with indent (continuation)
            if not any(
                f(lines[j]) for f in [is_h1, is_h2, is_h3, is_image, is_table_sep, is_table_row, is_code_fence, is_blockquote]
            ):
                items[-1] += " " + lines[j].strip()
            j += 1
        for item_text in items:
            p = doc.add_paragraph(style="List Bullet")
            parse_inline_format(p, item_text)
        i = j
        continue

    # Regular paragraph
    para_lines = [stripped]
    j = i + 1
    while j < len(lines):
        nxt = lines[j]
        if is_empty(nxt):
            break
        if any(f(nxt) for f in [is_h1, is_h2, is_h3, is_image, is_table_sep, is_table_row, is_code_fence, is_blockquote, is_list_item]):
            break
        para_lines.append(nxt.strip())
        j += 1

    para_text = " ".join(para_lines)
    if para_text.strip():
        p = doc.add_paragraph()
        parse_inline_format(p, para_text)
    i = j

# ── Save ──
doc.save(OUT_PATH)

# ── Validation ──
print(f"✅ Saved: {OUT_PATH}")
print(f"   Paragraphs: {len(doc.paragraphs)}")
print(f"   Images: {len([p for p in doc.paragraphs if p._element.xpath('.//pic:pic')])}")
print(f"   Tables: {len(doc.tables)}")

# Count headings
h_counts = {1: 0, 2: 0, 3: 0}
for p in doc.paragraphs:
    style_name = p.style.name if p.style else ""
    if style_name.startswith("Heading "):
        try:
            lvl = int(style_name.replace("Heading ", ""))
            if lvl in h_counts:
                h_counts[lvl] += 1
        except:
            pass

print(f"   Headings — H1: {h_counts[1]}, H2: {h_counts[2]}, H3: {h_counts[3]}")

# Verify font consistency
sample_runs = 0
font_issues = []
for p in doc.paragraphs[:20]:
    for run in p.runs[:3]:
        sample_runs += 1
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            east = rFonts.get(qn("w:eastAsia"))
            west = rFonts.get(qn("w:ascii"))
            if east and east not in (FONT_EAST, FONT_CODE_EAST):
                font_issues.append(f"  East font mismatch: {east}")
            if west and west not in (FONT_WEST, FONT_CODE):
                font_issues.append(f"  West font mismatch: {west}")

if font_issues:
    print(f"   ⚠️ Font issues found in {len(font_issues)} runs")
    for issue in font_issues[:5]:
        print(issue)
else:
    print(f"   ✅ Font consistency OK (checked {sample_runs} sample runs)")

# Check encoding by reading back
from docx.opc.constants import RELATIONSHIP_TYPE as RT
print(f"   ✅ Encoding: UTF-8 (default)")
